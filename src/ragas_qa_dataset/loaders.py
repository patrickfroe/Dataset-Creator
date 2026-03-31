from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class Document:
    text: str
    source_doc: str
    file_path: str
    file_type: str

    @property
    def source(self) -> str:
        """Backward-compatible alias."""
        return self.source_doc


def _read_txt_or_md(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("PDF support requires 'pypdf'.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("DOCX support requires 'python-docx'.") from exc

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()


def _normalize_file_types(file_types: tuple[str, ...]) -> set[str]:
    return {ft.lower().lstrip(".") for ft in file_types}


def _read_by_type(path: Path, file_type: str) -> str:
    if file_type in {"txt", "md"}:
        return _read_txt_or_md(path)
    if file_type == "pdf":
        return _read_pdf(path)
    if file_type == "docx":
        return _read_docx(path)
    raise ValueError(f"Unsupported file_type: {file_type}")


def load_local_documents(
    input_dir: Path,
    file_types: tuple[str, ...],
    max_docs: int | None = None,
) -> list[Document]:
    if not input_dir.exists():
        logger.error("Input directory does not exist: %s", input_dir)
        return []

    allowed_types = _normalize_file_types(file_types)
    documents: list[Document] = []

    for path in sorted(input_dir.rglob("*")):
        if not path.is_file():
            continue

        file_type = path.suffix.lower().lstrip(".")
        if file_type not in allowed_types:
            continue

        try:
            text = _read_by_type(path, file_type)
        except Exception as exc:  # explicit logging, continue with next file
            logger.exception("Failed to load '%s' (%s): %s", path, file_type, exc)
            continue

        if not text:
            logger.warning("Skipping empty document: %s", path)
            continue

        documents.append(
            Document(
                text=text,
                source_doc=path.name,
                file_path=str(path),
                file_type=file_type,
            )
        )

        if max_docs is not None and len(documents) >= max_docs:
            break

    return documents
